import docx

# Create document object 'd'
d = docx.Document('demo.docx')

# Document object has a list of paragraph objects
# print(d.paragraphs)

# Each paragraph has a text member variable
# print(d.paragraphs[1].text)

# Create a paragraph object 'p' from the second paragraph
p = d.paragraphs[1]

# Each paragraph has a list or run objects (a run is text until a change of style)
# print(p.runs)

# Each run object has a text member  variable
# print(p.runs[0].text)

# Each run object has 'bold', 'italic' and 'underline' member variables
# print(p.runs[1].bold)

# Change text and format using asigning
# p.runs[3].underline = True
# p.runs[3].text = 'italic and underlined.'

# d.save('demo2.docx')

# All paragraph and run objects have a style member variable
p.style = 'Title'
d.save('demo3.docx')
