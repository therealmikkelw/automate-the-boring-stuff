import docx

# Create blank docx document object
d = docx.Document()

# Using add_paragraph() method to ... add paragraphs
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')

# Using add_run() method to append text to a paragraph
p = d.paragraphs[0]
p.add_run('This is a new run.')

# Using asigning to change format of second run in 'p'
p.runs[1].bold = True
d.save('demo4.docx')

# When modifying an existing docx, open that and a new blank docx,
# and copy all paragraphs (or runs) to the blank docx while modifying.
